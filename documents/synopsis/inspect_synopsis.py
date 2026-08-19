from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


document = Document(Path(__file__).with_name("FF_180_reference.docx"))
cell = document.tables[4].cell(1, 0)
for index, paragraph in enumerate(cell.paragraphs):
    print(index, repr(paragraph.text), paragraph.style.name)

print("children")
for child in cell._tc.iterchildren():
    print(child.tag, child.find(qn("w:txbxContent")) is not None)

table_element = cell._tc.findall(qn("w:tbl"))[0]
print("literature table rows")
for row_index, row_element in enumerate(table_element.findall(qn("w:tr"))):
    cells = row_element.findall(qn("w:tc"))
    print(row_index, len(cells), ["".join(c.itertext())[:100] for c in cells])
