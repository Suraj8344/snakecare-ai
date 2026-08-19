from pathlib import Path
from zipfile import ZipFile

from docx import Document


DOCX_PATH = Path(__file__).with_name("FF_180_reference.docx")


document = Document(DOCX_PATH)
print(f"Tables: {len(document.tables)}")
for table_index, table in enumerate(document.tables):
    print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        print(
            row_index,
            [cell.text.replace("\n", " | ")[:160] for cell in row.cells],
            [id(cell._tc) for cell in row.cells],
        )

with ZipFile(DOCX_PATH) as archive:
    xml = archive.read("word/document.xml").decode("utf-8")
for needle in ("Author", "Objectives claimed", "Gaps Identified"):
    position = xml.find(needle)
    print(f"XML {needle!r}: {position}")
    print(xml[max(0, position - 500) : position + 1_500])
