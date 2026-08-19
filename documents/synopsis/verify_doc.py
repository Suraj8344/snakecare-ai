from __future__ import annotations

import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


DOCX_PATH = Path(__file__).with_name("SnakeCare_AI_Project_Synopsis_FF_180.docx")
EXPECTED = (
    "SnakeCare AI",
    "Artificial Intelligence, Digital Health, and Emergency Response Systems",
    "Project Synopsis:",
    "Background / Timeline:",
    "Gaps Identified:",
    "Objectives Framed Based on Gaps",
    "Problem Statement:",
    "Proposed Methodologies:",
    "Expected Outcomes:",
    "References (IEEE Format):",
    "10.2196/71378",
)


def main() -> int:
    document = Document(DOCX_PATH)
    if len(document.tables) != 8:
        raise AssertionError(f"Expected 8 top-level tables, found {len(document.tables)}")

    full_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    for phrase in EXPECTED:
        if phrase not in full_text:
            raise AssertionError(f"Missing expected phrase: {phrase}")

    synopsis_cell = document.tables[4].cell(1, 0)
    embedded_tables = synopsis_cell._tc.findall(qn("w:tbl"))
    if len(embedded_tables) != 1:
        raise AssertionError("Literature table was not preserved")
    rows = embedded_tables[0].findall(qn("w:tr"))
    if len(rows) != 4 or any(len(row.findall(qn("w:tc"))) != 5 for row in rows):
        raise AssertionError("Literature table must remain 4 rows by 5 columns")

    registration = document.tables[0]
    if "Department: Information Technology" not in registration.cell(3, 0).text:
        raise AssertionError("Department field changed unexpectedly")
    if "Group No. : 13" not in registration.cell(4, 2).text:
        raise AssertionError("Group number changed unexpectedly")
    if "Prof.Pravin.R.Futane" not in document.tables[2].cell(0, 0).text:
        raise AssertionError("Guide details changed unexpectedly")

    with ZipFile(DOCX_PATH) as archive:
        corrupt = archive.testzip()
        if corrupt:
            raise AssertionError(f"Corrupt package member: {corrupt}")
        required = {"word/document.xml", "word/styles.xml", "word/header1.xml"}
        missing = required.difference(archive.namelist())
        if missing:
            raise AssertionError(f"Missing package members: {sorted(missing)}")

    print(f"PASS: {DOCX_PATH.name}")
    print(f"Top-level tables: {len(document.tables)}")
    print(f"Synopsis paragraphs: {len(synopsis_cell.paragraphs)}")
    print("Literature table: 4 rows x 5 columns")
    print("Existing academic, student, and guide fields preserved")
    return 0


if __name__ == "__main__":
    sys.exit(main())
