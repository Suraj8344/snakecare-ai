from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


WORK_DIR = Path(__file__).resolve().parent
SOURCE = WORK_DIR / "FF_180_reference.docx"
OUTPUT = WORK_DIR / "SnakeCare_AI_Project_Synopsis_FF_180.docx"

TITLE = (
    "SnakeCare AI: Intelligent Snakebite Emergency Response, Dynamic Medical "
    "Passport, and Hospital Coordination Platform"
)
PROJECT_AREA = "Artificial Intelligence, Digital Health, and Emergency Response Systems"


def set_cell_text(cell, text: str, *, bold_label: str | None = None, size: float = 10.0):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    if bold_label and text.startswith(bold_label):
        label_run = paragraph.add_run(bold_label)
        label_run.bold = True
        value_run = paragraph.add_run(text[len(bold_label) :])
        runs = (label_run, value_run)
    else:
        runs = (paragraph.add_run(text),)
    for run in runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_run_font(paragraph, *, size: float = 10.0, bold: bool | None = None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def set_paragraph(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.0,
    left_indent: float = 0,
    space_after: float = 2,
):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Pt(left_indent)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_margins(cell, top=70, start=85, bottom=70, end=85):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def get_embedded_literature_table(cell):
    tables = cell._tc.findall(qn("w:tbl"))
    if len(tables) != 1:
        raise RuntimeError(f"Expected one embedded literature table, found {len(tables)}")
    return tables[0]


def xml_cells(row_element):
    return row_element.findall(qn("w:tc"))


def xml_cell_paragraph(cell_element):
    paragraph = cell_element.find(qn("w:p"))
    if paragraph is None:
        paragraph = OxmlElement("w:p")
        cell_element.append(paragraph)
    return paragraph


def set_xml_cell_text(cell_element, text: str, *, bold=False, size=7.5):
    for child in list(cell_element):
        if child.tag == qn("w:p"):
            cell_element.remove(child)
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "200")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    paragraph.append(p_pr)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(font_size)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    cell_element.append(paragraph)


def replace_registration_fields(document):
    identity = document.tables[0]
    set_cell_text(identity.cell(5, 0), f"Project Title: {TITLE}", bold_label="Project Title: ")
    set_cell_text(identity.cell(6, 0), f"Project Area: {PROJECT_AREA}", bold_label="Project Area: ")


def fill_literature_table(table_element):
    rows = table_element.findall(qn("w:tr"))
    header = [
        "Author / Year",
        "Objective",
        "Methodology",
        "Outcome",
        "Gap relevant to SnakeCare AI",
    ]
    for cell_element, text in zip(xml_cells(rows[0]), header, strict=True):
        set_xml_cell_text(cell_element, text, bold=True, size=7.5)

    evidence_rows = [
        [
            "GBD 2019 Snakebite Envenomation Collaborators (2022)",
            "Estimate global and national snakebite mortality from 1990–2019.",
            "GBD modelling using verbal-autopsy and vital-registration data.",
            "Estimated 63,400 global deaths in 2019; India carried the largest burden.",
            "Quantifies burden but does not provide a real-time patient-to-hospital workflow.",
        ],
        [
            "Ooms et al. (2021)",
            "Assess availability, affordability, and stock-outs of snakebite commodities.",
            "Cross-sectional survey of 133 Kenyan health facilities.",
            "Documented limited antivenom availability and frequent commodity stock-outs.",
            "Shows the need for authenticated, time-stamped inventory reporting and alerts.",
        ],
        [
            "Dash et al. (2025)",
            "Map digital-health interventions for snakebite management.",
            "PRISMA-ScR search; 237 records screened and 16 apps/interventions included.",
            "Apps support education, identification, mapping, and telemedicine, but evidence is uneven.",
            "Calls for offline access, local-system integration, reliability, and ethical safeguards.",
        ],
    ]
    for row_element, values in zip(rows[1:], evidence_rows, strict=True):
        for cell_element, text in zip(xml_cells(row_element), values, strict=True):
            set_xml_cell_text(cell_element, text, size=7.2)


def fill_synopsis(document):
    synopsis_cell = document.tables[4].cell(1, 0)
    set_margins(synopsis_cell)
    literature_table = get_embedded_literature_table(synopsis_cell)
    fill_literature_table(literature_table)

    paragraphs = synopsis_cell.paragraphs
    content = {
        0: ("Project Synopsis:", True, 11.0, 0),
        1: (
            "SnakeCare AI is a secure mobile and web platform that coordinates the time-critical snakebite "
            "journey from patient-reported symptoms to hospital readiness. It combines a dynamic medical "
            "passport, medical-report management, explainable triage decision support, location-aware hospital "
            "discovery, consented pre-alerts, and verified hospital resource operations. A QR workflow records "
            "antivenom-box usage only after authorization by verified hospital staff. The 112 voice-handoff "
            "feature is a local safety rehearsal and will not place calls or dispatch help until an approved "
            "ERSS integration and legal review exist.",
            False,
            9.5,
            0,
        ),
        2: ("", False, 9.5, 0),
        3: (f"Project Title: {TITLE}", True, 10.0, 0),
        4: ("", False, 9.5, 0),
        5: ("Introduction:", True, 10.5, 0),
        6: (
            "Background / Timeline: Snakebite envenoming is a WHO-recognized neglected tropical disease. "
            "WHO's 2019 strategy targets a 50% reduction in deaths and disability by 2030. Yet outcomes remain "
            "strongly affected by delayed transport, fragmented records, and uncertain antivenom availability. "
            "SnakeCare AI is developed as a 2026–27 academic prototype to connect patient, hospital, and "
            "authorized public-health workflows without replacing emergency services or clinical judgement.",
            False,
            9.5,
            10,
        ),
        7: (
            "Motivation: A patient or caregiver may not know which nearby facility can stabilize the patient, "
            "while the receiving hospital may lack advance information, recent records, or reliable stock data. "
            "The project aims to reduce avoidable coordination delay, preserve patient consent, and help hospitals "
            "locate nearby verified antivenom supply after the patient reaches the nearest appropriate facility.",
            False,
            9.5,
            10,
        ),
        8: (
            "Keywords: Snakebite emergency response; digital health; medical passport; explainable AI; hospital "
            "coordination; antivenom inventory; QR verification; consent; role-based access control.",
            False,
            9.5,
            10,
        ),
        9: ("", False, 9.5, 0),
        10: ("Literature Survey", True, 10.5, 0),
        11: (
            "The reviewed literature establishes the continuing burden of snakebite, recurring stock-visibility "
            "problems, and the promise and limitations of current digital interventions.",
            False,
            9.5,
            0,
        ),
        12: ("Table 1. Literature Survey of Existing Literature", True, 9.5, 0),
        13: ("", False, 9.0, 0),
        14: ("Gaps Identified:", True, 10.5, 0),
        15: (
            "1. Existing solutions are usually isolated tools for awareness, identification, or mapping rather "
            "than an end-to-end, consented patient-to-hospital workflow.",
            False,
            9.5,
            10,
        ),
        16: (
            "2. Hospital and antivenom information is often static, unverified, or not time-stamped; accountable "
            "staff approval and auditable QR stock updates are missing.",
            False,
            9.5,
            10,
        ),
        17: (
            "3. Many systems lack secure role verification, granular consent, explainable outputs, offline-safe "
            "essentials, provenance labels, and tested failure handling for emergency communication.",
            False,
            9.5,
            10,
        ),
        18: ("", False, 9.0, 0),
        19: ("Objectives Framed Based on Gaps", True, 10.5, 0),
        20: (
            "1. Build a modular cross-platform system for symptom intake, location capture, first-aid information, "
            "explainable severity support, nearby-hospital discovery, and consented pre-alerts.",
            False,
            9.5,
            10,
        ),
        21: (
            "2. Implement a secure dynamic medical passport and report repository with authentication, role-based "
            "authorization, explicit sharing consent, versioning, provenance, and audit logs.",
            False,
            9.5,
            10,
        ),
        22: (
            "3. Provide verified hospital operations for resource status and authorized QR-based antivenom usage, "
            "then evaluate reliability, usability, security, and emergency failure modes.",
            False,
            9.5,
            10,
        ),
        23: ("", False, 9.0, 0),
        24: ("Problem Statement:", True, 10.5, 0),
        25: (
            "Snakebite care is time-critical, but patients, caregivers, and hospitals frequently operate with "
            "fragmented medical information, uncertain resource availability, and weak coordination. A secure, "
            "auditable, and extensible platform is required to collect patient-reported emergency data, support "
            "safe routing to the nearest appropriate hospital, share only consented facts, and enable verified "
            "hospitals to coordinate antivenom resources without presenting AI output as diagnosis or treatment.",
            False,
            9.5,
            0,
        ),
        26: ("Proposed Methodologies:", True, 10.5, 0),
        27: (
            "a. Software and Hardware Requirements: Flutter/Dart with Material 3, Riverpod, GoRouter, Dio, and "
            "secure storage; FastAPI/Python with Pydantic, SQLAlchemy, Alembic, PostgreSQL, optional Redis, "
            "Firebase authentication, Gemini-assisted grounded intent classification, Docker, and CI/CD. Target "
            "hardware includes Android/web devices with GPS, camera, microphone, QR support, and a secure server.",
            False,
            9.2,
            10,
        ),
        28: (
            "b. Algorithms: schema validation; rule-based and explainable severity scoring; Haversine distance and "
            "capability-aware hospital ranking; OCR extraction with human confirmation; role/consent checks; QR "
            "inventory transactions; grounded AI intent classification that answers only from verified or "
            "patient-consented facts and returns unknown when evidence is absent.",
            False,
            9.2,
            10,
        ),
        29: (
            "c. Architecture: Flutter Client -> REST API -> FastAPI Service Layer -> Repository Layer -> "
            "PostgreSQL/Redis. External adapters provide Firebase identity, Gemini classification, maps/geocoding, "
            "and notifications. Hospital/government dashboards use server-controlled roles and immutable audit "
            "events. The emergency-handoff rehearsal is isolated from real 112/ERSS dispatch.",
            False,
            9.2,
            10,
        ),
        30: ("", False, 9.0, 0),
        31: ("Expected Outcomes:", True, 10.5, 0),
        32: (
            "A tested SnakeCare AI prototype that supports secure patient registration, dynamic medical passports, "
            "medical-report upload/OCR/search, explainable snakebite intake, nearest-hospital discovery, consented "
            "pre-alerts, verified hospital resource updates, and authorized QR antivenom accounting. Expected "
            "quality outcomes include reduced information-handoff time in simulation, traceable access and stock "
            "events, graceful handling of missing data/network failures, and an extensible architecture for other "
            "time-critical emergencies. Clinical effectiveness and real ERSS integration remain subjects for "
            "prospective validation and regulatory approval.",
            False,
            9.5,
            0,
        ),
        33: (
            "References (IEEE Format):\n"
            "[1] World Health Organization, Snakebite Envenoming: A Strategy for Prevention and Control. Geneva, "
            "Switzerland: WHO, 2019.\n"
            "[2] GBD 2019 Snakebite Envenomation Collaborators, “Global mortality of snakebite envenoming between "
            "1990 and 2019,” Nature Communications, vol. 13, Art. no. 6160, 2022, doi: "
            "10.1038/s41467-022-33627-9.\n"
            "[3] G. I. Ooms et al., “Availability, affordability and stock-outs of commodities for the treatment "
            "of snakebite in Kenya,” PLOS Neglected Tropical Diseases, vol. 15, no. 8, Art. no. e0009702, 2021, "
            "doi: 10.1371/journal.pntd.0009702.\n"
            "[4] A. Dash, S. Kerketta, G. Mallick, J. Menon, S. Kanungo, and S. Pati, “Digital Health Intervention "
            "in Snakebite Management: Scoping Review,” Journal of Medical Internet Research, vol. 27, Art. no. "
            "e71378, 2025, doi: 10.2196/71378.",
            False,
            8.8,
            0,
        ),
    }
    for index, (text, bold, size, indent) in content.items():
        set_paragraph(
            paragraphs[index],
            text,
            bold=bold,
            size=size,
            left_indent=indent,
            space_after=2 if text else 0,
        )


def main():
    document = Document(SOURCE)
    replace_registration_fields(document)
    fill_synopsis(document)
    document.core_properties.title = "SnakeCare AI Project Synopsis"
    document.core_properties.subject = "FF No. 180 Project Registration and Synopsis"
    document.core_properties.keywords = (
        "SnakeCare AI, snakebite, digital health, medical passport, hospital coordination"
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
